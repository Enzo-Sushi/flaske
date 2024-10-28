import os
from flask import Flask, render_template, request, send_file, abort
from docx import Document
import pdfplumber

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = os.path.join(os.getcwd(), 'uploads')  # Define o caminho absoluto

# Certifique-se de que a pasta uploads existe
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert_pdf_to_word():
    pdf_file = request.files['pdf_file']
    
    # Verifique e defina um caminho de arquivo válido
    pdf_filename = pdf_file.filename
    pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
    pdf_file.save(pdf_path)
    print(f"Arquivo PDF salvo em: {pdf_path}")

    # Defina o caminho do arquivo Word com o mesmo nome do PDF, mas com extensão .docx
    word_filename = pdf_filename.rsplit('.', 1)[0] + '.docx'
    doc_path = os.path.join(app.config['UPLOAD_FOLDER'], word_filename)

    try:
        # Criar o documento Word e iniciar o processo de escrita
        doc = Document()
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    doc.add_heading(f'Texto da página {i + 1}', level=1)
                    doc.add_paragraph(text)
                else:
                    print(f"A página {i + 1} está vazia ou não pôde ser extraída.")

        # Salve o documento Word
        doc.save(doc_path)
        print(f"Documento Word salvo em: {doc_path}")

        # Verifique se o arquivo foi realmente criado
        if not os.path.exists(doc_path):
            print("Erro: O arquivo Word não foi criado.")
            abort(500, description="Erro na criação do arquivo Word.")
        else:
            print("Arquivo Word criado com sucesso.")

    except Exception as e:
        print(f"Erro ao converter o PDF: {e}")
        abort(500, description="Erro ao converter o PDF para Word.")

    # Verifique o tamanho do arquivo antes de enviar para download
    try:
        file_size = os.path.getsize(doc_path)
        print(f"Tamanho do arquivo gerado: {file_size} bytes")  # Confirmação do tamanho do arquivo
    except Exception as e:
        print(f"Erro ao acessar o arquivo Word para enviar: {e}")
        abort(500, description="Erro ao acessar o arquivo Word para download.")

    # Enviar o arquivo para download
    return send_file(doc_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
